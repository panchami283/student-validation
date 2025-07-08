from docx import Document
import os

def generate_docx(data):
    if not os.path.exists("output"):
        os.makedirs("output")

    doc = Document()
    doc.add_heading("Admission Form - Verified", 0)

    for key, value in data.items():
        doc.add_paragraph(f"{key}: {value}")
    
    doc.add_paragraph("Status: ✅ Verified")

    full_name = "_".join(filter(None, [data.get('first_name'), data.get('middle_name'), data.get('last_name')])).replace(" ", "_")
    output_path = os.path.join("output", f"admission_{full_name}.docx")

    doc.save(output_path)
